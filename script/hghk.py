import json

from docx import Document
from docx.oxml.ns import qn
import re


month_en = ['january', 'february', 'march', 'april', 'may', 'june', 'july', 'august', 'september', 'october', 'november', 'december']

def get_cell_fill_color(cell):
    """
    Retrieves the fill color (shading) of a table cell.
    Returns the hex color string if found, otherwise None.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shading_elm = tc_pr.find(qn('w:shd'))
    if shading_elm is not None:
        fill_color = shading_elm.get(qn('w:fill'))
        return fill_color
    return None

def extract_table_to_dataframe(table):
    days_dict = {
        
    }
    for i in range(1, table.rows.__len__()):
        for cell in table.rows[i].cells:
            if 'ЛИЧНЫЙ ДЕНЬ'.lower() in cell.text.lower():
                text = cell.text.replace('\t', ' ')
                texts = text.split(' ')[0]

                days = texts.split('\n')
                day = int(days[0])
                personal_day = int(days[1])
                bg = get_cell_fill_color(cell)
                if bg == 'F59A00':
                    color = 'orange'
                elif bg == '50A5A0':
                    color = 'green'
                elif bg == 'FA0701':
                    color = 'red'
                else:
                    color = None
                days_dict[str(day)] = {
                    'day': day,
                    'personal_day': personal_day,
                    'day_by_color': color
                }
    return days_dict

def loop_month(tables, n,year,target_dict:dict):
    i = 0
    for table in tables:
        print(n, i)
        days = extract_table_to_dataframe(table)
        for day in days:
            target_dict[year][f'member_{n}']['calendar'][month_en[i]][day] = days[day]
        i += 1
        

with open('../db/y2025.json', 'r', encoding='utf-8') as f:
    
    target = json.load(f)
    for i in range(1,6):
        document = Document(r"C:\Users\Admin\Downloads\Год {} S (2025 КС).docx".format(i))
        tables = document.tables
        print(len(tables))
        loop_month(tables, i, '2025', target)
    for i in range(6,10):   
        document = Document(r"C:\Users\Admin\Downloads\Год {} S (2025 СК).docx".format(i))
        tables = document.tables
        print(len(tables))
        loop_month(tables, i, '2025', target)
    # for i in range(1, 10):
    #     target['2025'][f'member_{i}']['member'] = i
    with open('../db/y2025.json', 'w', encoding='utf-8') as ff:
        json.dump(target, ff, ensure_ascii=False, indent=4)
        
        


